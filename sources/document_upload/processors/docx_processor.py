"""
DOCX document processor.
"""

import os
from datetime import datetime
from typing import List
from docx import Document

from sources.base.interfaces import DocumentProcessor, SourceResult


class DOCXProcessor(DocumentProcessor):
    """Processor for DOCX documents."""
    
    def supports_format(self, file_extension: str) -> bool:
        """Check if this processor supports the given file format."""
        return file_extension.lower() == '.docx'
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return ['.docx']
    
    def process_document(self, file_path: str, filename: str) -> SourceResult:
        """Process a DOCX document and return structured content."""
        try:
            # Extract text from DOCX
            doc = Document(file_path)
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Get file metadata
            file_stats = os.stat(file_path)
            created_date = datetime.fromtimestamp(file_stats.st_ctime)
            updated_date = datetime.fromtimestamp(file_stats.st_mtime)
            
            # Try to get document properties
            title = filename
            author = None
            try:
                core_props = doc.core_properties
                if core_props.title:
                    title = core_props.title
                if core_props.author:
                    author = core_props.author
                if core_props.created:
                    created_date = core_props.created
                if core_props.modified:
                    updated_date = core_props.modified
            except:
                pass  # Ignore metadata extraction errors
            
            return SourceResult(
                content=full_text.strip(),
                metadata={
                    'file_path': file_path,
                    'file_size': file_stats.st_size,
                    'paragraph_count': len(doc.paragraphs),
                    'processor': 'DOCXProcessor'
                },
                source_type='document_upload',
                source_id=f"docx_{filename}_{file_stats.st_mtime}",
                title=title,
                author=author,
                created_date=created_date,
                updated_date=updated_date,
                tags=['docx', 'document', 'word']
            )
            
        except Exception as e:
            raise RuntimeError(f"Failed to process DOCX {filename}: {str(e)}") 